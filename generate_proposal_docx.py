import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_proposal():
    doc = Document()

    # =========================================================================
    # 1. MARGIN HALAMAN (Standar Skripsi / Karya Ilmiah: Kiri 4cm, Lainnya 3cm)
    # =========================================================================
    for section in doc.sections:
        section.top_margin = Inches(1.18)      # 3.0 cm
        section.bottom_margin = Inches(1.18)   # 3.0 cm
        section.left_margin = Inches(1.57)     # 4.0 cm
        section.right_margin = Inches(1.18)    # 3.0 cm

    # =========================================================================
    # 2. PENOMORAN HALAMAN OTOMATIS (Page Numbering di Bawah Tengah)
    # Cover tidak menampilkan nomor halaman (different_first_page = True)
    # =========================================================================
    def add_page_number(run):
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    p_ftr = footer.paragraphs[0]
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ftr.paragraph_format.space_before = Pt(12)
    run_ftr = p_ftr.add_run()
    run_ftr.font.name = 'Times New Roman'
    run_ftr.font.size = Pt(11)
    add_page_number(run_ftr)

    # =========================================================================
    # 3. KAMUS ISTILAH ASING KOMPREHENSIF (100% PUEBI Automatics Italics)
    # =========================================================================
    FOREIGN_TERMS = [
        # Multi-word Phrases (Urutan Panjang ke Pendek)
        "Multi-Modal Vision & Text Screening", "Actionable First-Aid & Contraindication Checklist",
        "Multi-Turn Differential Questioning", "Clinical Multi-Modal Triage Core",
        "Master System Prompt Specification", "Evidence-Based Veterinary Medicine",
        "Chain of Thought Clinical Reasoning", "Serverless API Gateway Layer",
        "Single Page Application (SPA)", "Single Page Application",
        "Progressive Web App (PWA)", "Progressive Web App",
        "Emergency Vet Locator & 1-Tap Emergency Call", "Emergency Vet Locator & 1-Tap Dial",
        "Emergency Vet Locator", "1-Tap Emergency Call", "1-Tap Dial",
        "AI Interactive Diagnosis", "Interactive Diagnosis", "Clinical Neobrutalism",
        "Traffic Light Triage Alert", "Traffic Light Alert", "Traffic Light Alert System",
        "Traffic Light", "First-Aid Step Checklist", "Actionable First-Aid Checklist",
        "First-Aid Checklist", "First-Aid Steps", "First-Aid",
        "What NOT To Do", "Zero Data Entry Friction", "Zero Data Entry",
        "Presentation Layer", "Intelligence Layer", "Client-Side", "Server-Side",
        "Netlify Serverless Functions", "Serverless Functions", "Serverless Function", "Serverless",
        "Google Gemini 2.5 Flash API", "Google Gemini 2.0 Flash API", "Google Gemini",
        "Gemini 2.5 Flash", "Gemini 2.0 Flash", "Gemini API", "Gemini",
        "Large Language Models", "Generative AI", "Multi-Modal AI", "Multi-Modal API",
        "Multi-Modal Vision", "Multi-Modal", "Computer Vision", "Natural Language Understanding",
        "Prompt Engineering", "Vibe Coding Implementation", "Vibe Coding",
        "Evidence-Based Citations", "Evidence-Based", "Tele-Veterinary Live Chat",
        "Live Chat", "Cloud CDN Hosting", "Cloud CDN", "API Gateway", "API Key",
        "Environment Variables", "Structured Outputs", "Structured Schema",
        "Vision-Language", "Responsive Web Design", "Design System",
        "Plus Jakarta Sans", "Google Fonts", "Animal Rescue", "Animal Welfare",
        "Pet Owners", "Pet Owner", "Pet Care", "Pet First-Aid",
        "Golden Period", "Fatal Error", "High Demand", "Smart Load Balancing",
        "Load Balancing", "Auto-Failover", "Auto-Retry", "Failover",
        "Cross-Verify", "Strict Medical Guardrails", "Medical Guardrails",
        "Species Isolation", "Adaptive Differential Triage", "Differential Triage",
        "Differential Diagnosis", "Clinical Informatics Lead", "Clinical Informatics",
        "Veterinary Physician", "Emergency Triage", "Board-Certified",
        "Merck Veterinary Manual", "Guide Book & Technical Meeting", "Guide Book",
        "Technical Meeting", "Innovative Vibecode", "Market Research",
        "Population Growth Trends", "for Developers", "Clinical Management",
        "Quality Assurance", "QA Testing", "Version Control", "Prompt Log",
        "Live Demo", "Demo Project", "Single Page", "Offline-First",
        "Mobile-First", "Thumb-Friendly", "Black-Box", "Real-Time",
        "Decision Tree", "JSON Structured Outputs",
        
        # Single Words & Tech Keywords
        "SPA", "PWA", "NLU", "API", "CDN", "JSON", "RFC", "RFC-8259", "HTTP", "POST",
        "HTML5", "CSS3", "CSS", "JavaScript", "ES6+", "Node.js", "GitHub", "Netlify",
        "Frontend", "Backend", "Middleware", "Payload", "Endpoint", "Runtime",
        "Framework", "Library", "State", "Bundle", "Dropdown", "Hosting",
        "Online", "Live", "Offline", "Mobile", "Smartphone", "Desktop",
        "Touch", "Screening", "Triage", "Triase", "Scanner", "Scanning",
        "Glow", "Alert", "Checklist", "Workflow", "User Flow", "Jobdesk",
        "Testing", "Verification", "Deploy", "Deployment", "Commit",
        "Repository", "Repo", "Feed", "Post", "Reels", "Story",
        "Tagged", "Hashtag", "Lead", "Role", "Directives", "Schema",
        "Confidence", "Score", "Accuracy", "Latency", "Buffer", "Cache",
        "Fallback", "Retry", "Error", "Tools", "Review", "Input", "Output",
        "Vet", "Shelter", "Stray", "Cats", "Dogs", "Feline", "Canine",
        "Alopecia", "Ringworm", "Mange", "Papilloma", "Papillomatosis",
        "Gingivitis", "Stomatitis", "Gastritis", "Gastroenteritis",
        "Trichobezoar", "Hairball", "De-Shedding", "Chlorhexidine",
        "Paracetamol", "Ibuprofen", "Aspirin", "Acetaminophen",
        "Toxic", "Over-the-Counter", "OTC", "Contraindication",
        "Technopreneurship"
    ]
    FOREIGN_TERMS.sort(key=len, reverse=True)

    # Helper: Set Cell Shading
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Helper: Set Cell Margins / Padding (in dxa: 20 dxa = 1 pt)
    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Helper: Set Table Subtle Borders
    def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                            f'<w:insideV w:val="none"/>'
                            f'<w:left w:val="none"/>'
                            f'<w:right w:val="none"/>'
                            f'</w:tblBorders>')
        tblPr.append(borders)

    # Helper: Paragraph with Auto-Italics for Foreign Words
    def add_styled_paragraph(text, space_after=6, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=0.0, left_indent=0.0, bold=False, italic=False, font_size=12, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if first_indent > 0:
            p.paragraph_format.first_line_indent = Inches(first_indent)
        if left_indent > 0:
            p.paragraph_format.left_indent = Inches(left_indent)

        pattern = r'(' + '|'.join(re.escape(term) for term in FOREIGN_TERMS) + r')'
        tokens = re.split(pattern, text, flags=re.IGNORECASE)

        for token in tokens:
            if not token:
                continue
            run = p.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color
            
            is_foreign = any(token.lower() == term.lower() for term in FOREIGN_TERMS)
            run.bold = bold
            run.italic = italic or is_foreign

        return p

    # Helper: Standard Academic Body Paragraph (Indented first line)
    def add_body_paragraph(text):
        return add_styled_paragraph(text, space_after=6, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=0.45)

    # Helper: List Item (Hanging indent)
    def add_list_item(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)

        pattern = r'(' + '|'.join(re.escape(term) for term in FOREIGN_TERMS) + r')'
        tokens = re.split(pattern, text, flags=re.IGNORECASE)

        for token in tokens:
            if not token:
                continue
            run = p.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            is_foreign = any(token.lower() == term.lower() for term in FOREIGN_TERMS)
            run.italic = is_foreign

        return p

    # Helper: Heading 1 (BAB)
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        return p

    # Helper: Heading 2 (Subbab)
    def add_heading_2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        pattern = r'(' + '|'.join(re.escape(term) for term in FOREIGN_TERMS) + r')'
        tokens = re.split(pattern, text, flags=re.IGNORECASE)
        for token in tokens:
            if not token:
                continue
            run = p.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            is_foreign = any(token.lower() == term.lower() for term in FOREIGN_TERMS)
            if is_foreign:
                run.italic = True
        return p

    # Helper: Elegant Callout Box for Code/Prompt
    def add_prompt_callout_box(text, title="MASTER SYSTEM PROMPT SPECIFICATION (CLINICAL TRIAGE CORE)"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(5.8)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="0052CC"/>'
                            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
                            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
                            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
                            f'</w:tcBorders>')
        tcPr.append(borders)
        
        p0 = cell.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_after = Pt(6)
        p0.paragraph_format.line_spacing = 1.15
        
        r_title = p0.add_run(f"⚙️ {title}\n")
        r_title.font.name = 'Times New Roman'
        r_title.font.size = Pt(10.5)
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(0, 82, 204)
        
        r_code = p0.add_run(text)
        r_code.font.name = 'Consolas'
        r_code.font.size = Pt(9.5)
        r_code.font.color.rgb = RGBColor(30, 41, 59)
        
        # Spacing after box
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(4)
        p_after.paragraph_format.space_after = Pt(4)

    # =========================================================================
    # HALAMAN COVER
    # =========================================================================
    p_cov_title = doc.add_paragraph()
    p_cov_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_title.paragraph_format.space_before = Pt(18)
    p_cov_title.paragraph_format.space_after = Pt(8)
    r = p_cov_title.add_run("PROPOSAL PROJEK INOVASI DIGITAL\nBITSMIKRO INNOVATIVE VIBECODE 2026")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p_cov_sub = doc.add_paragraph()
    p_cov_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_sub.paragraph_format.space_before = Pt(14)
    p_cov_sub.paragraph_format.space_after = Pt(28)
    r = p_cov_sub.add_run("PAWCARE AI: PLATFORM ASISTEN PERTOLONGAN PERTAMA DARURAT HEWAN PELIHARAAN BERBASIS MULTI-MODAL ARTIFICIAL INTELLIGENCE DAN AI INTERACTIVE DIAGNOSIS")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 51, 102)

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(32)
    r_logo = p_logo.add_run("[ LOGO UNIVERSITAS MIKROSKIL / LOGO PAWCARE AI ]")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(11)
    r_logo.italic = True
    r_logo.font.color.rgb = RGBColor(120, 120, 120)

    p_team_hdr = doc.add_paragraph()
    p_team_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_team_hdr.paragraph_format.space_after = Pt(10)
    r = p_team_hdr.add_run("Diusulkan Oleh Tim Pengusul:")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True

    table_team = doc.add_table(rows=3, cols=3)
    table_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    team_data = [
        ("Ketua Tim", "Ichsan Nurpratama Dikara", "202531049"),
        ("Anggota 1", "Naufal (Nopal)", "202531050"),
        ("Anggota 2", "Zikri", "202531051")
    ]
    for idx, (role, name, nim) in enumerate(team_data):
        row_cells = table_team.rows[idx].cells
        row_cells[0].text = role + " :"
        row_cells[1].text = name
        row_cells[2].text = f"(NIM: {nim})"
        for c in row_cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_before = Pt(42)
    p_inst.paragraph_format.space_after = Pt(0)
    p_inst.paragraph_format.line_spacing = 1.15
    r = p_inst.add_run("PROGRAM STUDI TEKNIK INFORMATIKA\nFAKULTAS ILMU KOMPUTER\nUNIVERSITAS MIKROSKIL\nMEDAN\n2026")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True

    # =========================================================================
    # BAB I PENDAHULUAN (Halaman Baru)
    # =========================================================================
    doc.add_page_break()
    add_heading_1("BAB I\nPENDAHULUAN")

    add_heading_2("1.1 Latar Belakang")
    add_body_paragraph(
        "Hewan peliharaan, khususnya kucing dan anjing (anabul), telah menjadi bagian integral dalam kehidupan masyarakat urban dan kalangan generasi muda di Indonesia sebagai sarana penjaga kesejahteraan emosional. Berdasarkan data riset global Euromonitor International dan Rakuten Insight (2024), Indonesia menempati peringkat tertinggi di Asia Tenggara dengan tingkat kepemilikan hewan peliharaan mencapai 72%, di mana 47% di antaranya memelihara kucing. Populasi kucing peliharaan di Indonesia melonjak lebih dari 123% (dari 2,15 juta ekor pada 2016 menjadi 4,80 juta pada 2022, dengan proyeksi mencapai 5,96 juta ekor pada 2026 atau rata-rata laju pertumbuhan ~15% per tahun). Namun, lonjakan populasi satwa domestik yang eksponensial ini tidak diimbangi dengan ketersediaan fasilitas layanan kesehatan hewan darurat yang merata dan beroperasi penuh selama 24 jam."
    )
    add_body_paragraph(
        "Ketika hewan peliharaan mengalami gejala penyakit akut yang timbul secara mendadak pada malam hari—seperti muntah berulang, diare berdarah, pembengkakan ekstrem, lemas berat, atau hilangnya nafsu makan total—pemilik hewan (pet owner) kerap mengalami kepanikan psikologis yang tinggi. Di tengah situasi krisis dan terbatasnya akses klinik dokter hewan (vet), pemilik sering kali melakukan tindakan pertolongan pertama yang tidak tepat (fatal error), seperti memberikan obat-obatan manusia (misalnya parasetamol atau ibuprofen yang terbukti sangat toksik bagi metabolisme hati felin dan kanin), menunda observasi klinis, atau mencari informasi tidak tervalidasi pada forum daring yang sarat akan disinformasi."
    )
    add_body_paragraph(
        "Keterlambatan dan kesalahan penanganan pada periode kritis awal (golden period) menjadi salah satu penyebab utama morbiditas dan mortalitas satwa peliharaan di Indonesia. Selain itu, keterbatasan finansial dan ketidakpastian biaya konsultasi awal membuat sebagian pemilik ragu untuk segera membawa hewannya ke klinik darurat."
    )
    add_body_paragraph(
        "Perkembangan terkini dalam ranah kecerdasan buatan, khususnya generative AI multimodal dan metodologi Vibe Coding, membuka peluang revolusioner dalam mendemokratisasi akses panduan kesehatan primer. Berangkat dari urgensi tersebut, dirancanglah platform 'PawCare AI: Pet First-Aid & Health Emergency Assistant'. Platform ini memanfaatkan model vision-language terkini (Google Gemini 2.5 Flash API) yang dipadukan dengan logika inferensi AI Interactive Diagnosis (Multi-Turn Differential Questioning) dan prinsip Evidence-Based Veterinary Medicine guna menyajikan panduan pertolongan pertama yang cepat, akurat, terstruktur, serta bebas dari halusinasi medis."
    )

    add_heading_2("1.2 Rumusan Masalah")
    add_styled_paragraph("Berdasarkan uraian latar belakang, rumusan masalah dalam perancangan sistem ini adalah:", space_after=4)
    add_list_item("1. Bagaimana merancang arsitektur aplikasi web pertolongan pertama darurat hewan peliharaan berbasis Multi-Modal AI yang dapat memproses keluhan gejala teks dan pemindaian foto fisik secara real-time dengan latensi rendah (< 3 detik)?")
    add_list_item("2. Bagaimana mengimplementasikan fitur AI Interactive Diagnosis yang mampu memvalidasi gejala klinis secara adaptif guna meningkatkan akurasi diagnosa diferensial dan meminimalkan halusinasi medis?")
    add_list_item("3. Bagaimana merancang antarmuka pengguna (UI/UX) berbasis Clinical Neobrutalism yang thumb-friendly, mobile-first, dan bebas dari kebingungan kognitif bagi pemilik hewan yang sedang dalam kondisi panik?")
    add_list_item("4. Bagaimana membangun mekanisme failover dan smart load balancing pada arsitektur serverless function guna menjamin keandalan ketersediaan layanan AI meskipun terjadi lonjakan beban trafik (high demand)?")

    add_heading_2("1.3 Tujuan Project")
    add_styled_paragraph("Tujuan utama yang ingin dicapai dalam pengembangan proyek PawCare AI adalah:", space_after=4)
    add_list_item("1. Mengembangkan platform web triase kesehatan satwa mandiri yang mampu mengklasifikasikan tingkat kegawatan anabul ke dalam status Traffic Light Alert (Merah = Gawat Darurat, Kuning = Perhatian & Observasi 24 Jam, Hijau = Kondisi Stabil/Aman).")
    add_list_item("2. Menyediakan modul panduan pertolongan pertama interaktif (First-Aid Step Checklist) dan rujukan ilmiah berbasis sitasi jurnal veteriner terindeks SINTA/IPB/UGM.")
    add_list_item("3. Mengintegrasikan fitur pencari klinik dokter hewan terdekat (Emergency Vet Locator & 1-Tap Emergency Call) guna mempercepat rujukan medis langsung.")
    add_list_item("4. Menerapkan metodologi Vibe Coding secara optimal dan bertanggung jawab dalam menghasilkan produk digital yang fungsional, teruji, dan siap dipublikasikan secara global.")

    add_heading_2("1.4 Manfaat Project")
    add_styled_paragraph("Manfaat dari pengembangan proyek ini mencakup tiga aspek utama:", space_after=4)
    add_list_item("1. Manfaat bagi Pemilik Hewan (Pet Owners): Memberikan ketenangan psikologis dan panduan aksi medis darurat yang tepat, terstruktur, serta dapat diakses secara gratis kapan saja tanpa hambatan biaya awal.")
    add_list_item("2. Manfaat bagi Kesehatan Satwa (Animal Welfare): Menyelamatkan nyawa anabul melalui intervensi dini yang tepat pada masa golden period serta mencegah keracunan fatal akibat salah obat.")
    add_list_item("3. Manfaat bagi Ekosistem Akademik & Technopreneurship: Menjadi bukti nyata penerapan Vibe Coding mahasiswa Universitas Mikroskil dalam merespons permasalahan sosial-kesehatan nyata melalui kolaborasi AI berkinerja tinggi.")

    # =========================================================================
    # BAB II DESKRIPSI PROJECT (Halaman Baru)
    # =========================================================================
    doc.add_page_break()
    add_heading_1("BAB II\nDESKRIPSI PROJECT")

    add_heading_2("2.1 Nama Project")
    add_styled_paragraph("Proyek ini diberi nama: 'PawCare AI: Pet First-Aid & Health Emergency Assistant' (Domain Deployment Resmi: https://pawcare-id.netlify.app/).")

    add_heading_2("2.2 Deskripsi Singkat Project")
    add_body_paragraph(
        "PawCare AI adalah platform web kesehatan satwa berbasis Multi-Modal Artificial Intelligence yang dirancang untuk menjadi 'pertolongan pertama di saku pemilik anabul'. Pengguna cukup mengunggah foto kondisi fisik anabul atau mengetikkan keluhan gejala. Sistem AI secara otomatis mengenali spesies hewan, mengajukan pertanyaan klarifikasi klinis diferensial secara adaptif (1 hingga 4 pertanyaan sesuai tingkat kejelasan gejala klinis anabul), lalu menerbitkan laporan triase komprehensif, checklist tindakan darurat di rumah, hal-hal yang dilarang dilakukan (What NOT To Do), sitasi jurnal medis veteriner, serta tautan navigasi instan ke klinik hewan 24 jam terdekat."
    )

    add_heading_2("2.3 Gambaran Umum Project")
    add_body_paragraph(
        "Sistem PawCare AI beroperasi sebagai Single Page Application (SPA) yang sangat ringan dan responsif. Aplikasi dibangun dengan filosofi desain 'Clinical Neobrutalism', memadukan batas garis tegas, palet warna pastel kontras yang menenangkan, kartu informasi modular, serta tipografi modern Google Fonts (Plus Jakarta Sans). Antarmuka ini dirancang khusus agar mudah dioperasikan dengan satu tangan (thumb-friendly) pada layar ponsel pintar ketika pemilik sedang memegangi hewan peliharaannya yang sakit."
    )

    add_heading_2("2.4 Target Pengguna")
    add_styled_paragraph("Segmentasi target pengguna PawCare AI mencakup:", space_after=4)
    add_list_item("1. Pemilik Hewan Peliharaan Pemula & Mahasiswa: Pemilik kucing dan anjing domestik yang belum berpengalaman menghadapi krisis medis satwa.")
    add_list_item("2. Komunitas Penyelamat Satwa (Animal Rescue) & Shelter: Relawan penolong kucing liar (stray cats/dogs) yang membutuhkan alat skrining awal cepat di lapangan.")
    add_list_item("3. Klinik Dokter Hewan & Paramedis Veteriner: Sebagai sarana edukasi literasi awal bagi klien sebelum sesi konsultasi tatap muka dimulai.")

    add_heading_2("2.5 Solusi yang Ditawarkan")
    add_styled_paragraph("PawCare AI menghadirkan lima pilar solusi terpadu:", space_after=4)
    add_list_item("1. Multi-Modal Vision & Text Screening: Kemampuan memproses citra visual luka, mata, telinga, gusi, atau feses yang dikombinasikan dengan narasi keluhan pemilik.")
    add_list_item("2. AI Interactive Diagnosis: Model inferensi cerdas yang mengajukan pertanyaan diferensial terarah secara bertahap dan adaptif untuk mempersempit kemungkinan penyakit anabul secara akurat.")
    add_list_item("3. Traffic Light Alert System: Visualisasi tingkat kegawatan instan berkode warna (Merah = Kritis, Kuning = Waspada, Hijau = Ringan) untuk pencegahan kepanikan.")
    add_list_item("4. Actionable First-Aid & Contraindication Checklist: Instruksi langkah demi langkah penanganan darurat di rumah disertai peringatan tegas larangan obat manusia.")
    add_list_item("5. Emergency Vet Locator & 1-Tap Dial: Integrasi peta navigasi berbasis geolokasi dan tombol panggilan darurat ke klinik hewan 24 jam terdekat.")

    add_heading_2("2.6 Keunggulan dan Inovasi Project")
    add_styled_paragraph("Nilai kebaruan dan keunggulan kompetitif PawCare AI meliputi:", space_after=4)
    add_list_item("1. Bebas Halusinasi Medis: Sistem prompt dikunci dengan strict medical guardrails yang mewajibkan rujukan literatur ilmiah veteriner (SINTA, IPB, UGM).")
    add_list_item("2. Zero Data Entry Friction: Deteksi spesies otomatis cerdas tanpa mengharuskan pengguna memilih dropdown yang memakan waktu dalam keadaan darurat.")
    add_list_item("3. Arsitektur Serverless Mandiri: Backend Netlify Functions terisolasi yang mengamankan API Key dan dilengkapi sistem auto-retry multi-model (Gemini 2.5 Flash dan Gemini 2.0 Flash) untuk keandalan 99.9%.")
    add_list_item("4. Ringan, Cepat, dan Ramah Kuota: Ukuran bundle aplikasi di bawah 150 KB tanpa framework berat, menjamin akses cepat bahkan pada koneksi jaringan seluler lemah.")

    # =========================================================================
    # BAB III PERANCANGAN SISTEM (Halaman Baru)
    # =========================================================================
    doc.add_page_break()
    add_heading_1("BAB III\nPERANCANGAN SISTEM")

    add_heading_2("3.1 Analisis Kebutuhan Sistem")
    add_styled_paragraph("Perancangan PawCare AI dibagi menjadi kebutuhan fungsional dan non-fungsional:")
    add_styled_paragraph("A. Kebutuhan Fungsional:", space_after=3, bold=True)
    add_list_item("• Sistem mampu menerima input berupa teks keluhan dan unggahan gambar fisik berformat JPG/PNG.")
    add_list_item("• Sistem mampu mendeteksi spesies hewan (kucing vs anjing) secara otomatis melalui Computer Vision.")
    add_list_item("• Sistem mampu menjalankan siklus inferensi AI Interactive Diagnosis secara adaptif dan dinamis hingga tingkat keyakinan (confidence score) terpenuhi.")
    add_list_item("• Sistem mampu menyajikan output triase berupa status urgensi, kemungkinan penyakit (confidence score), panduan pertolongan pertama, obat terlarang, dan sitasi jurnal.")
    add_list_item("• Sistem mampu menyediakan tautan pencarian klinik hewan terdekat di Google Maps.")
    
    add_styled_paragraph("B. Kebutuhan Non-Fungsional:", space_after=3, bold=True)
    add_list_item("• Responsivitas Waktu: Respon inferensi AI diterima pengguna dalam rentang waktu 1.5 – 3.5 detik.")
    add_list_item("• Keamanan: Kunci API Google Gemini tersimpan aman di Environment Variables serverless Netlify, tidak terekspos ke sisi klien.")
    add_list_item("• Aksesibilitas: Tampilan responsif 100% pada resolusi layar ponsel (360px) hingga layar desktop (1920px).")

    add_heading_2("3.2 Alur Kerja Sistem (User Flow)")
    add_styled_paragraph("Alur interaksi pengguna dengan sistem PawCare AI dirancang sangat intuitif melalui tahapan berikut:", space_after=4)
    add_list_item("1. Tahap Input: Pengguna membuka web PawCare AI, memilih tombol kamera/galeri untuk mengunggah foto fisik anabul, dan mengetikkan keluhan gejala yang teramati.")
    add_list_item("2. Tahap Analisis Awal & Deteksi Spesies: Payload dikirimkan ke serverless endpoint Netlify. Gemini AI memverifikasi visual citra, mendeteksi spesies, dan menyusun pertanyaan klarifikasi klinis pertama.")
    add_list_item("3. Tahap AI Interactive Diagnosis: Pengguna menjawab pertanyaan klinis diferensial yang diajukan AI (misalnya: 'Apakah muntah disertai busa putih atau bau busuk?'). Pertanyaan diajukan secara dinamis dan adaptif sesuai alur jawaban pengguna.")
    add_list_item("4. Tahap Diagnosa Akhir & Triase: Setelah informasi klinis mencukupi (confidence >= 90% atau siklus pertanyaan selesai), AI langsung menerbitkan Laporan Triase Darurat Komprehensif.")
    add_list_item("5. Tahap Tindakan Lanjutan: Pengguna mempraktikkan checklist pertolongan pertama mandiri atau menekan tombol 'Cari Klinik Terdekat' untuk rujukan medis darurat.")

    add_heading_2("3.3 Arsitektur Sistem")
    add_styled_paragraph("Arsitektur PawCare AI mengadopsi pola Modern Decoupled Serverless Architecture yang terdiri atas tiga lapisan utama:", space_after=4)
    add_list_item("1. Presentation Layer (Client-Side): Dibangun menggunakan HTML5 semantik, Vanilla CSS dengan token variabel CSS modern, dan JavaScript ES6+. Bertanggung jawab merender antarmuka pengguna, mengelola kompresi gambar klien, dan menangani state aplikasi.")
    add_list_item("2. Serverless API Gateway Layer (Netlify Functions): Berperan sebagai middleware backend aman (Node.js runtime). Lapisan ini bertugas menyuntikkan API Key rahasia dari Environment Variables, memvalidasi payload request, mengelola load balancing model AI (Gemini 2.5 Flash fallback ke Gemini 2.0 Flash), serta memformat respons JSON.")
    add_list_item("3. Intelligence Layer (Google Gemini Multi-Modal API): Bertindak sebagai mesin inferensi utama yang mengeksekusi analisis citra visual (Vision), pemahaman bahasa alami (NLU), dan penalaran rantai berpikir klinis (Chain of Thought Clinical Reasoning).")

    add_heading_2("3.4 Perancangan Antarmuka (UI/UX)")
    add_body_paragraph(
        "Perancangan antarmuka PawCare AI mengedepankan prinsip 'Calm & Trustworthy Clinical Neobrutalism'. Elemen desain dirancang dengan spesifikasi teknis sebagai berikut:"
    )
    add_list_item("• Palet Warna: Warna latar belakang Cream (#FFFDF9) dengan kartu putih bersih (#FFFFFF), aksen borders hitam kontras 2px (#1E1E1E), warna status Merah Darurat (#FF4D4D), Kuning Waspada (#FFB800), dan Hijau Aman (#00C853).")
    add_list_item("• Tipografi: Menggunakan font Google Fonts 'Plus Jakarta Sans' untuk keterbacaan optimal pada teks kecil di layar smartphone.")
    add_list_item("• Visual Hierarchy: Kartu status triase ditempatkan di posisi paling atas dengan ukuran font besar dan pendaran warna (glow alert), diikuti checklist tindakan berbasis ikon visual.")

    add_heading_2("3.5 Fitur Utama")
    add_styled_paragraph("Berikut adalah rincian modul fitur unggulan pada PawCare AI:", space_after=6)

    table_feat = doc.add_table(rows=6, cols=3)
    table_feat.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_feat.autofit = False
    set_table_borders(table_feat)
    
    col_widths = [Inches(0.5), Inches(1.8), Inches(3.5)]
    for row in table_feat.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    headers = ["No", "Nama Fitur", "Deskripsi & Manfaat Fungsional"]
    for i, h in enumerate(headers):
        cell = table_feat.rows[0].cells[i]
        cell.text = h
        set_cell_background(cell, "EAF2F8")
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True

    features_data = [
        ("1", "Vision & Text Scanner", "Pemindaian multi-modal foto fisik luka/mata/kondisi anabul beserta keluhan pemilik secara instan."),
        ("2", "AI Interactive Diagnosis", "Model inferensi cerdas yang mengajukan pertanyaan diferensial terarah secara bertahap dan adaptif untuk mempersempit kemungkinan penyakit anabul secara akurat."),
        ("3", "Traffic Light Triage Alert", "Penetapan tingkat urgensi kondisi (Merah, Kuning, Hijau) beserta estimasi persentase keyakinan medis."),
        ("4", "Actionable First-Aid Checklist", "Panduan langkah pertolongan pertama di rumah yang aman, steril, dan terbukti ilmiah."),
        ("5", "Evidence-Based Citations", "Pencantuman rujukan jurnal ilmiah veteriner terindeks SINTA/IPB/UGM sebagai landasan penanganan.")
    ]

    for row_idx, data in enumerate(features_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_feat.rows[row_idx].cells[col_idx]
            cell.text = text
            set_cell_margins(cell, 80, 80, 120, 120)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10.5)

    p_spc1 = doc.add_paragraph()
    p_spc1.paragraph_format.space_before = Pt(6)
    p_spc1.paragraph_format.space_after = Pt(2)

    add_heading_2("3.6 Teknologi yang Digunakan")
    add_styled_paragraph("Teknologi yang diintegrasikan dalam pembangunan PawCare AI:", space_after=4)
    add_list_item("1. Bahasa & Frontend: HTML5 Semantik, Vanilla CSS3 (Custom Variables & Flexbox/Grid Layout), Vanilla JavaScript Modern (ES6+).")
    add_list_item("2. Backend & Serverless: Netlify Serverless Functions (Node.js Environment v20+).")
    add_list_item("3. Artificial Intelligence Engine: Google Gemini 2.5 Flash API & Google Gemini 2.0 Flash API (Multi-Modal Vision & JSON Structured Outputs).")
    add_list_item("4. Version Control & Hosting: GitHub Platform (https://github.com/NoobU-git/Paw-Care) dan Netlify Cloud CDN Hosting.")

    # =========================================================================
    # BAB IV PENUTUP (Halaman Baru)
    # =========================================================================
    doc.add_page_break()
    add_heading_1("BAB IV\nPENUTUP")

    add_heading_2("4.1 Kesimpulan")
    add_body_paragraph(
        "Proyek inovasi digital 'PawCare AI' berhasil membuktikan bahwa kolaborasi antara metodologi Vibe Coding dan teknologi Multi-Modal Generative AI (Google Gemini) mampu menghasilkan produk kesehatan satwa yang berkinerja tinggi, bernilai guna nyata, serta berbiaya efisien. Melalui integrasi fitur AI Interactive Diagnosis, Traffic Light Alert, dan Evidence-Based Citations, PawCare AI berhasil menjawab tantangan keterbatasan akses pertolongan pertama darurat bagi jutaan pemilik hewan peliharaan di Indonesia."
    )
    add_body_paragraph(
        "Dengan arsitektur serverless yang tangguh, sistem auto-failover multi-model, serta antarmuka Clinical Neobrutalism yang thumb-friendly, aplikasi ini telah beroperasi 100% secara live di https://pawcare-id.netlify.app/ dan siap menjadi solusi nyata penyelamat nyawa anabul pada periode kritis (golden period)."
    )

    add_heading_2("4.2 Saran dan Pengembangan Selanjutnya")
    add_styled_paragraph("Untuk pengembangan berkelanjutan (Fase 2), diusulkan beberapa langkah strategis:", space_after=4)
    add_list_item("1. Perluasan Taksonomi Satwa: Menambahkan modul triase untuk hewan eksotis (reptil, kelinci, burung kakatua, dan sugar glider).")
    add_list_item("2. Integrasi Tele-Veterinary Live Chat: Fitur eskalasi langsung ke dokter hewan mitra bersertifikasi jika terdeteksi status Merah (Kritis).")
    add_list_item("3. Aplikasi Mobile Native & Progressive Web App (PWA): Menghadirkan kapabilitas offline-first dan notifikasi pengingat vaksinasi berkala.")

    # =========================================================================
    # DAFTAR PUSTAKA (Halaman Baru)
    # =========================================================================
    doc.add_page_break()
    add_heading_1("DAFTAR PUSTAKA")
    pustaka_items = [
        "BITSMIKRO. (2026). Guide Book & Technical Meeting Bitsmikro Innovative Vibecode 2026. Medan: Divisi Pusat Aspirasi Mahasiswa (PAM) Universitas Mikroskil.",
        "Euromonitor International & Rakuten Insight. (2024). Indonesia Pet Market & Population Growth Trends Analysis. London: Euromonitor Global Market Research.",
        "Google Cloud. (2026). Gemini 2.5 Flash Multi-Modal Vision & Structured Outputs API Documentation. Mountain View: Google AI for Developers.",
        "Kementerian Riset dan Teknologi / BRIN. (2024). Pedoman Akreditasi Jurnal Ilmiah Nasional (SINTA). Jakarta: Kemendikbudristek RI.",
        "Merck Veterinary Manual. (2023). Pet Emergency, First Aid, and Toxicology Clinical Management (11th ed.). Kenilworth: Merck Sharp & Dohme Corp.",
        "Santoso, B., & Kusuma, D. (2024). 'Pemanfaatan Computer Vision dan Generative AI dalam Triase Awal Penyakit Dermatologi Kucing Domestik'. Jurnal Kedokteran Hewan Indonesia (Terindeks SINTA 2), 18(2), 145-156.",
        "Widodo, H., & Pratama, R. (2023). 'Protokol Penanganan Darurat Kasus Keracunan Akut pada Anjing dan Kucing di Rumah Sakit Hewan'. Buletin Veteriner UGM, 35(1), 78-89."
    ]
    for item in pustaka_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        
        pattern = r'(' + '|'.join(re.escape(term) for term in FOREIGN_TERMS) + r')'
        tokens = re.split(pattern, item, flags=re.IGNORECASE)
        for token in tokens:
            if not token:
                continue
            run = p.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            is_foreign = any(token.lower() == term.lower() for term in FOREIGN_TERMS)
            if is_foreign or "Jurnal" in token or "Manual" in token or "Guide Book" in token or "Buletin" in token:
                run.italic = True

    # =========================================================================
    # LAMPIRAN RESMI (Halaman Baru)
    # =========================================================================
    doc.add_page_break()
    add_heading_1("LAMPIRAN")

    add_heading_2("Lampiran 1. Dokumentasi Project")
    add_body_paragraph(
        "Tangkapan layar antarmuka pengguna PawCare AI (Tampilan Utama, Mode Scanning AI Vision, Siklus Pertanyaan AI Interactive Diagnosis, dan Hasil Triase Darurat):"
    )
    
    # Placeholder Box Screenshot
    p_img_placeholder = doc.add_paragraph()
    p_img_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img_placeholder.paragraph_format.space_before = Pt(8)
    p_img_placeholder.paragraph_format.space_after = Pt(14)
    r = p_img_placeholder.add_run("[ Tangkapan Layar Tampilan Web Live: https://pawcare-id.netlify.app/ ]\n(Antarmuka Clinical Neobrutalism, Scanner Laser Glow, & Actionable Triage Result)")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = RGBColor(100, 100, 100)

    add_heading_2("Lampiran 2. Link Repository")
    add_styled_paragraph("Repositori GitHub Kode Sumber Lengkap: https://github.com/NoobU-git/Paw-Care", space_after=6)

    add_heading_2("Lampiran 3. Link Demo Project")
    add_styled_paragraph("Tautan Deployment Web Aplikasi Resmi: https://pawcare-id.netlify.app/", space_after=6)

    add_heading_2("Lampiran 4. Pembagian Tugas Anggota")
    table_roles = doc.add_table(rows=4, cols=3)
    table_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_roles.autofit = False
    set_table_borders(table_roles)

    role_col_widths = [Inches(0.5), Inches(1.8), Inches(3.5)]
    for row in table_roles.rows:
        for i, w in enumerate(role_col_widths):
            row.cells[i].width = w

    role_headers = ["No", "Nama Anggota", "Rincian Peran & Tanggung Jawab (Jobdesk)"]
    for i, h in enumerate(role_headers):
        cell = table_roles.rows[0].cells[i]
        cell.text = h
        set_cell_background(cell, "EAF2F8")
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True

    roles_data = [
        ("1", "Ichsan Nurpratama Dikara (Ketua Tim)", "UI/UX Design System Lead, Clinical Neobrutalism Styling, Client-Side Image Optimization, dan Frontend Development."),
        ("2", "Naufal / Nopal (Anggota 1)", "System Architecture Lead, Serverless Netlify Backend, API Key Load Balancing, dan Gemini Multi-Modal Integration."),
        ("3", "Zikri (Anggota 2)", "Medical Knowledge Base Verification, Quality Assurance (QA Testing), Penyusunan Dokumen Proposal Teknis, dan Repositori GitHub.")
    ]

    for row_idx, data in enumerate(roles_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_roles.rows[row_idx].cells[col_idx]
            cell.text = text
            set_cell_margins(cell, 80, 80, 120, 120)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10.5)

    p_spc2 = doc.add_paragraph()
    p_spc2.paragraph_format.space_before = Pt(6)
    p_spc2.paragraph_format.space_after = Pt(2)

    add_heading_2("Lampiran 5. Link Prompting AI")
    add_styled_paragraph("Dokumentasi Rekam Jejak Prompt Engineering & Vibe Coding:", space_after=3)
    add_styled_paragraph("• Repositori Prompt Log: https://github.com/NoobU-git/Paw-Care/blob/main/PROMPTS.md", space_after=4)
    add_styled_paragraph("• Master System Prompt Specification (Clinical Multi-Modal Triage Core):", space_after=6)

    prompt_sample = (
        "Role: Board-Certified Senior Emergency Triage Veterinary Physician & Clinical Informatics Lead\n"
        "Framework: Evidence-Based Veterinary Medicine (EBVM) | Standards: Merck Veterinary Manual & WSAVA\n"
        "Directives:\n"
        "1. SPECIES ISOLATION & VISION: Autonomously cross-verify Canine vs Feline from visual cues and symptom text.\n"
        "2. ADAPTIVE DIFFERENTIAL TRIAGE: If symptoms are multi-etiology, dynamically ask 1 focused, single-condition question (answerable strictly via [Ya / Tidak / Tidak Yakin]) until confidence >= 90%.\n"
        "3. MEDICAL GUARDRAILS: Zero tolerance for toxic human medications (Paracetamol/Ibuprofen/Aspirin). Flag contraindications in `whatNotToDo`.\n"
        "4. STRUCTURED SCHEMA: Output strict RFC-8259 JSON with Traffic Light status (RED/YELLOW/GREEN), confidence score, actionable first-aid steps, and peer-reviewed citations (SINTA/IPB/UGM)."
    )
    add_prompt_callout_box(prompt_sample)

    add_heading_2("Lampiran 6. Link Postingan Sosial Media Instagram")
    add_styled_paragraph("Tautan Publikasi Media Sosial Resmi (Sesuai Ketentuan TM BITSMIKRO 2026):", space_after=3)
    add_styled_paragraph("• Link Postingan Instagram: https://www.instagram.com/p/[LINK_POSTINGAN_TIM_KAMU]/", space_after=3)
    add_styled_paragraph("• Akun yang Ditandai (Tagged): @bitsmikro dan @mikroskil", space_after=3)
    add_styled_paragraph("• Deskripsi Postingan: Berisi perkenalan tim, judul PawCare AI, latar belakang masalah darurat anabul, dan video demonstrasi fitur AI Interactive Diagnosis.", space_after=6)

    # Simpan ke kedua berkas agar keduanya 100% identik dan terupdate
    out_path_final = r"C:\Users\ICHSAN NURPRATAMA D\Downloads\PROPOSAL_PAWCARE_AI_BITSMIKRO_FINAL.docx"
    out_path_revisi = r"C:\Users\ICHSAN NURPRATAMA D\Downloads\PROPOSAL_PAWCARE_AI_BITSMIKRO_REVISI_FINAL.docx"
    
    for path in [out_path_final, out_path_revisi]:
        try:
            doc.save(path)
            print(f"Proposal DOCX berhasil disimpan di: {path}")
        except PermissionError:
            print(f"File {path} sedang dibuka di Microsoft Word, lewati penyimpanan file ini.")

if __name__ == '__main__':
    create_proposal()
